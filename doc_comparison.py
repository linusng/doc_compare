"""
Document Comparison Tool — Full-Context, No RAG
================================================
Compares a base PDF against a Word letter of offer by feeding both
documents in a single LLM call. No chunking, no vector store, no
retrieval — the model sees every clause in full context simultaneously.

Why no RAG:
  - Both documents fit comfortably within a 128K context window
  - Chunking destroys cross-clause context (e.g. conditions defined
    in clause 2 that modify clause 7 would be invisible to a chunk)
  - Embedding similarity is a poor proxy for legal materiality
  - A single holistic pass also catches omissions (clauses present in
    the base but entirely absent from the offer)

Architecture:
  - PyMuPDF  : PDF extraction with page tracking
  - python-docx : Word extraction with page-break detection
  - Pydantic v2 : all data models + structured LLM output schema
  - openai   : direct OpenAI-compatible client (no LangChain needed)
  - tiktoken : pre-flight token count guard

Usage:
    python doc_comparison.py --base base.pdf --compare offer.docx
    python doc_comparison.py --base base.pdf --compare offer.docx \\
        --output report.md --base-url https://your-endpoint/v1

Requirements:
    pip install pymupdf python-docx openai pydantic tiktoken

    export OPENAI_API_KEY="your-key"
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path
from typing import Optional

import anthropic
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from langchain_ollama import ChatOllama
from openai import OpenAI
from pydantic import BaseModel, Field, model_validator


# ─────────────────────────────────────────────
# Pydantic models
# ─────────────────────────────────────────────

class Paragraph(BaseModel):
    """A single extracted paragraph with its source metadata."""
    text: str
    page: int                      # 1-based
    doc_name: str
    section: str = "Introduction"


class PagedDocument(BaseModel):
    """Full extracted document, paragraphs in reading order."""
    doc_name: str
    paragraphs: list[Paragraph]

    def as_prompt_block(self) -> str:
        """
        Render the document as a numbered, page-annotated text block
        suitable for inclusion in an LLM prompt.

        Format per paragraph:
            [p.3 | §Employment Terms] The employee shall receive...
        """
        lines: list[str] = []
        for i, p in enumerate(self.paragraphs, start=1):
            lines.append(f"[{i}. p.{p.page} | §{p.section}] {p.text}")
        return "\n\n".join(lines)


class DeviationItem(BaseModel):
    """
    One deviation found by the LLM between the two documents.
    All fields are validated by Pydantic — deviation is always bool.
    """
    item_no: int = Field(description="Sequential deviation number, starting at 1.")
    section: str = Field(description="Short overarching category label inferred from the subject matter of the compared clauses (e.g. 'Interest Rate', 'Security', 'Repayment'). Must be concise (2–4 words), not a verbatim heading from either document.")
    base_page: int = Field(description="Page number in the base PDF where the clause appears.")
    compare_page: int = Field(description="Page number in the offer document where the clause appears.")
    base_paragraph: str = Field(description="The relevant clause text from the base document.")
    compare_paragraph: str = Field(description="The corresponding clause text from the offer document.")
    deviation: bool = Field(description="True if a material deviation exists, False otherwise.")
    comments: Optional[str] = Field(
        default=None,
        description=(
            "Concise one-sentence explanation of the deviation (max 30 words). "
            "Must be None when deviation is False."
        ),
    )

    @model_validator(mode="after")
    def comment_required_when_deviation(self) -> "DeviationItem":
        if self.deviation and not self.comments:
            raise ValueError("comments must be provided when deviation is True")
        if not self.deviation and self.comments is not None:
            # Silently clear comments on non-deviations to keep output clean
            object.__setattr__(self, "comments", None)
        return self


class ComparisonResult(BaseModel):
    """Top-level structured output returned by the LLM."""
    deviations: list[DeviationItem] = Field(
        description="List of ALL clauses examined, including those where deviation=False."
    )

    @property
    def deviation_count(self) -> int:
        return sum(1 for d in self.deviations if d.deviation)

    @property
    def flagged(self) -> list[DeviationItem]:
        return [d for d in self.deviations if d.deviation]


class ComparisonReport(BaseModel):
    """Final report wrapping the result with document metadata."""
    base_doc: str
    compare_doc: str
    total_tokens_used: int = 0
    result: ComparisonResult


# ─────────────────────────────────────────────
# Document loaders
# ─────────────────────────────────────────────

def load_pdf(path: str) -> PagedDocument:
    """
    Extract paragraphs from a PDF using PyMuPDF's block-level API.

    get_text("blocks", sort=True) returns text blocks in reading order.
    Each block is a tuple: (x0, y0, x1, y1, text, block_no, block_type)
    block_type 0 = text, 1 = image — we skip images.
    """
    doc_name = Path(path).name
    paragraphs: list[Paragraph] = []

    with fitz.open(path) as doc:
        for page_num, page in enumerate(doc, start=1):
            for block in page.get_text("blocks", sort=True):
                if block[6] != 0:
                    continue
                text = block[4].strip().replace("\n", " ")
                if len(text) > 30:
                    paragraphs.append(Paragraph(
                        text=text,
                        page=page_num,
                        doc_name=doc_name,
                    ))

    return PagedDocument(doc_name=doc_name, paragraphs=paragraphs)


_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _page_breaks_in_paragraph(para) -> int:
    """
    Count page breaks inside a paragraph's XML, covering two sources:

    1. w:lastRenderedPageBreak — inserted automatically by Word at soft/auto
       page-break positions when the file is saved after rendering. This is
       what the old `run.contains_page_break` missed entirely.

    2. w:br w:type="page" — explicit hard page breaks (Ctrl+Enter).
    """
    count = 0
    for elem in para._p.iter():
        local = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if local == "lastRenderedPageBreak":
            count += 1
        elif local == "br" and elem.get(f"{{{_W}}}type") == "page":
            count += 1
    return count


def load_docx(path: str) -> PagedDocument:
    """
    Extract paragraphs from a Word document.
    Page numbers are tracked by scanning each paragraph's raw XML for both
    w:lastRenderedPageBreak (soft breaks inserted by Word on save) and
    w:br[@w:type='page'] (explicit hard breaks).
    """
    doc_name = Path(path).name
    paragraphs: list[Paragraph] = []
    page = 1

    for para in DocxDocument(path).paragraphs:
        page += _page_breaks_in_paragraph(para)
        text = para.text.strip()
        if text and len(text) > 30:
            paragraphs.append(Paragraph(
                text=text,
                page=page,
                doc_name=doc_name,
            ))

    return PagedDocument(doc_name=doc_name, paragraphs=paragraphs)


# ─────────────────────────────────────────────
# Section inference
# ─────────────────────────────────────────────

def _looks_like_heading(text: str) -> bool:
    words = text.split()
    if len(words) > 12 or text.endswith("."):
        return False
    return (
        text.istitle()
        or text.isupper()
        or bool(re.match(r"^\d+[\.\)]\s+\w", text))
    )


def infer_sections(doc: PagedDocument) -> PagedDocument:
    """
    Tag each paragraph with the most recently seen heading-like text.
    Uses model_copy to avoid mutating Pydantic model instances.
    """
    current = "Introduction"
    enriched: list[Paragraph] = []
    for p in doc.paragraphs:
        if _looks_like_heading(p.text):
            current = p.text
        enriched.append(p.model_copy(update={"section": current}))
    return doc.model_copy(update={"paragraphs": enriched})


# ─────────────────────────────────────────────
# Token guard
# ─────────────────────────────────────────────

def estimate_tokens(text: str, model: str = "") -> int:
    """
    Conservative character-based token estimate (≈ 4 chars/token for English).
    Used only as a pre-flight context-limit guard — the 85% safety margin
    absorbs estimation error across different model tokenisers.
    """
    return max(1, len(text) // 4)


# ─────────────────────────────────────────────
# Prompt builder
# ─────────────────────────────────────────────

SYSTEM_PROMPT = """You are a senior legal document reviewer specialising in \
contract and offer-letter compliance.

Your task is to perform a thorough clause-by-clause comparison between a \
BASE DOCUMENT (the reference standard) and a LETTER OF OFFER (the document \
under review).

Rules:
1. Compare every substantive clause in the base document against the \
corresponding clause in the offer document.
2. A material deviation exists when meaning, obligation, right, amount, \
date, timeline, or condition differs — even subtly.
3. Minor cosmetic rephrasing that preserves identical meaning is NOT a \
deviation.
4. If a clause present in the base document is entirely ABSENT from the \
offer, that is a deviation. Set compare_page to 0 and note the omission.
5. If a clause exists in the offer but NOT in the base document, that is \
also a deviation. Set base_page to 0 and note the addition.
6. Return ALL clauses examined, including those where deviation=False.
7. Populate every field accurately, especially page numbers.
8. The "section" field must be a SHORT category label (2–4 words) describing the subject matter of the deviation — e.g. "Interest Rate", "DSCR Covenant", "Security", "Repayment Terms". Never copy clause text into this field.

Respond strictly with a JSON object matching this schema — no markdown \
fences, no commentary outside the JSON:

{
  "deviations": [
    {
      "item_no": <int>,
      "section": "<short category label, 2-4 words, inferred from the subject matter — e.g. 'Interest Rate', 'Security', 'Repayment'>",
      "base_page": <int>,
      "compare_page": <int>,
      "base_paragraph": "<string>",
      "compare_paragraph": "<string>",
      "deviation": <bool>,
      "comments": "<string | null>"
    }
  ]
}"""


def build_user_prompt(base_doc: PagedDocument, compare_doc: PagedDocument) -> str:
    return (
        f"## BASE DOCUMENT — {base_doc.doc_name}\n\n"
        f"{base_doc.as_prompt_block()}\n\n"
        f"{'─' * 60}\n\n"
        f"## LETTER OF OFFER — {compare_doc.doc_name}\n\n"
        f"{compare_doc.as_prompt_block()}"
    )


# ─────────────────────────────────────────────
# LLM call
# ─────────────────────────────────────────────

def _token_guard(full_prompt: str, model: str, context_limit: int) -> None:
    estimated = estimate_tokens(full_prompt, model)
    print(f"   → Estimated prompt tokens : {estimated:,}")
    print(f"   → Context limit           : {context_limit:,}")
    if estimated > int(context_limit * 0.85):
        raise RuntimeError(
            f"Estimated prompt ({estimated:,} tokens) exceeds 85% of the "
            f"context limit ({context_limit:,}). Consider a model with a "
            f"larger context window or trim the documents."
        )


def run_comparison_openai(
    client: OpenAI,
    model: str,
    base_doc: PagedDocument,
    compare_doc: PagedDocument,
    context_limit: int,
) -> ComparisonReport:
    """OpenAI / OpenAI-compatible endpoint path (uses response_format JSON)."""
    user_prompt = build_user_prompt(base_doc, compare_doc)
    _token_guard(SYSTEM_PROMPT + "\n\n" + user_prompt, model, context_limit)

    response = client.chat.completions.create(
        model=model,
        temperature=0,
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user",   "content": user_prompt},
        ],
    )

    raw_json = response.choices[0].message.content
    tokens_used = response.usage.total_tokens if response.usage else 0

    try:
        parsed = json.loads(raw_json)
        result = ComparisonResult.model_validate(parsed)
    except (json.JSONDecodeError, Exception) as exc:
        raise RuntimeError(
            f"LLM returned invalid JSON or schema mismatch: {exc}\n"
            f"Raw response:\n{raw_json[:500]}"
        ) from exc

    return ComparisonReport(
        base_doc=base_doc.doc_name,
        compare_doc=compare_doc.doc_name,
        total_tokens_used=tokens_used,
        result=result,
    )


def _fix_section_labels(deviations: list[DeviationItem], llm: ChatOllama) -> list[DeviationItem]:
    """
    Post-processing pass: if any section label looks like copied clause text
    (> 50 chars), ask the model to re-label all deviations in one shot.
    """
    if all(len(d.section) <= 50 for d in deviations):
        return deviations

    items = "\n".join(
        f"{d.item_no}. base: {d.base_paragraph[:120]} | compare: {d.compare_paragraph[:120]}"
        for d in deviations
    )
    prompt = (
        "For each numbered deviation below, respond with ONLY a JSON array of short "
        "category labels (2–4 words each), in the same order. "
        "Examples: [\"Interest Rate\", \"DSCR Covenant\", \"Security\"]\n\n"
        + items
    )
    raw = llm.invoke([("human", prompt)]).content.strip()

    # Extract the JSON array from the response
    match = re.search(r"\[.*\]", raw, re.DOTALL)
    if not match:
        return deviations
    try:
        labels: list[str] = json.loads(match.group())
    except json.JSONDecodeError:
        return deviations

    # Flatten in case the model returns nested lists e.g. [["Interest Rate"], ...]
    flat: list[str] = []
    for item in labels:
        flat.append(item[0] if isinstance(item, list) else item)

    fixed = []
    for i, d in enumerate(deviations):
        label = str(flat[i]).strip() if i < len(flat) else d.section
        fixed.append(d.model_copy(update={"section": label}))
    return fixed


def run_comparison_ollama(
    model: str,
    base_doc: PagedDocument,
    compare_doc: PagedDocument,
    context_limit: int,
    base_url: str,
) -> ComparisonReport:
    """Ollama local path — uses LangChain with_structured_output to enforce schema."""
    user_prompt = build_user_prompt(base_doc, compare_doc)
    _token_guard(SYSTEM_PROMPT + "\n\n" + user_prompt, model, context_limit)

    llm = ChatOllama(model=model, temperature=0, base_url=base_url, num_ctx=context_limit)
    structured_llm = llm.with_structured_output(ComparisonResult)

    result: ComparisonResult = structured_llm.invoke([
        ("system", SYSTEM_PROMPT),
        ("human",  user_prompt),
    ])

    fixed_deviations = _fix_section_labels(result.deviations, llm)
    result = ComparisonResult(deviations=fixed_deviations)

    return ComparisonReport(
        base_doc=base_doc.doc_name,
        compare_doc=compare_doc.doc_name,
        total_tokens_used=0,   # Ollama does not expose token usage via LangChain
        result=result,
    )


def run_comparison_anthropic(
    model: str,
    base_doc: PagedDocument,
    compare_doc: PagedDocument,
    context_limit: int,
    api_key: Optional[str] = None,
) -> ComparisonReport:
    """Anthropic Claude path — uses client.messages.parse() for native Pydantic structured output."""
    user_prompt = build_user_prompt(base_doc, compare_doc)
    _token_guard(SYSTEM_PROMPT + "\n\n" + user_prompt, model, context_limit)

    client = anthropic.Anthropic(api_key=api_key or os.environ.get("ANTHROPIC_API_KEY", ""))

    with client.messages.stream(
        model=model,
        max_tokens=8192,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_prompt}],
        output_config={"format": {"type": "json_schema", "schema": ComparisonResult.model_json_schema()}},
    ) as stream:
        final = stream.get_final_message()

    raw_json = next(b.text for b in final.content if b.type == "text")
    tokens_used = final.usage.input_tokens + final.usage.output_tokens

    try:
        result = ComparisonResult.model_validate_json(raw_json)
    except Exception as exc:
        raise RuntimeError(
            f"Anthropic returned invalid JSON or schema mismatch: {exc}\n"
            f"Raw response:\n{raw_json[:500]}"
        ) from exc

    return ComparisonReport(
        base_doc=base_doc.doc_name,
        compare_doc=compare_doc.doc_name,
        total_tokens_used=tokens_used,
        result=result,
    )


# ─────────────────────────────────────────────
# Markdown renderer
# ─────────────────────────────────────────────

def _truncate(text: str, max_chars: int = 300) -> str:
    text = text.replace("\n", " ").strip()
    return text if len(text) <= max_chars else text[:max_chars].rstrip() + "…"


def render_markdown(report: ComparisonReport) -> str:
    all_items = report.result.deviations

    if not all_items:
        body = "_No clauses examined._\n"
    else:
        header = (
            f"| # | Section "
            f"| Source PDF `{report.base_doc}` Page "
            f"| Compared Doc `{report.compare_doc}` Page "
            f"| Base Document Paragraph "
            f"| Compared Document Paragraph "
            f"| Deviation "
            f"| Comments |\n"
            "|---|---|---|---|---|---|---|---|\n"
        )
        rows = []
        for d in all_items:
            base_pg = f"p.{d.base_page}"    if d.base_page    else "—"
            cmp_pg  = f"p.{d.compare_page}" if d.compare_page else "—"
            rows.append(
                f"| {d.item_no} "
                f"| {d.section} "
                f"| {base_pg} "
                f"| {cmp_pg} "
                f"| {_truncate(d.base_paragraph)} "
                f"| {_truncate(d.compare_paragraph)} "
                f"| {'Yes' if d.deviation else 'No'} "
                f"| {d.comments or '—'} |\n"
            )
        body = header + "".join(rows)

    token_line = (
        f"**Total tokens used:** {report.total_tokens_used:,}  \n"
        if report.total_tokens_used else ""
    )

    return (
        f"# Document Comparison Report\n\n"
        f"**Base document:** `{report.base_doc}`  \n"
        f"**Compared document:** `{report.compare_doc}`  \n"
        f"**Material deviations found:** {report.result.deviation_count}  \n"
        f"{token_line}"
        f"\n---\n\n"
        f"## Deviation Table\n\n"
        f"{body}\n"
        f"---\n"
        f"*Generated by doc_comparison.py — full-context, no RAG*\n"
    )


# ─────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Full-context document comparison — no RAG, single LLM call."
    )
    parser.add_argument("--base",          required=True,  help="Path to base PDF document")
    parser.add_argument("--compare",       required=True,  help="Path to Word (.docx) offer document")
    parser.add_argument("--output",        default="/output/comparison_report.md", help="Output Markdown file")
    parser.add_argument("--provider",      default="auto", choices=["auto", "openai", "ollama", "anthropic"],
                                           help="LLM backend: auto (detect from --base-url), openai, ollama, or anthropic")
    parser.add_argument("--model",         default="gpt-oss-120b", help="Model name (use claude-opus-4-6 for Anthropic)")
    parser.add_argument("--base-url",      default=None,   help="OpenAI-compatible API base URL")
    parser.add_argument("--api-key",       default=None,   help="API key (falls back to OPENAI_API_KEY / ANTHROPIC_API_KEY env var)")
    parser.add_argument("--context-limit", default=131072, type=int, help="Model context window in tokens (default: 131072)")
    args = parser.parse_args()

    # ── Load documents ───────────────────────────────────────────────────
    print(f"📄 Loading base PDF:     {args.base}")
    base_doc = infer_sections(load_pdf(args.base))
    print(f"   → {len(base_doc.paragraphs)} paragraphs extracted")

    print(f"📝 Loading compare DOCX: {args.compare}")
    compare_doc = infer_sections(load_docx(args.compare))
    print(f"   → {len(compare_doc.paragraphs)} paragraphs extracted")

    # ── Resolve provider ─────────────────────────────────────────────────
    base_url = args.base_url or ""
    provider = args.provider
    if provider == "auto":
        provider = "ollama" if ("localhost" in base_url or "127.0.0.1" in base_url) else "openai"

    print(f"\n🤖 Running full-context comparison with {args.model}…")
    if provider == "ollama":
        print(f"   → Backend: Ollama ({base_url})")
        report = run_comparison_ollama(
            model=args.model,
            base_doc=base_doc,
            compare_doc=compare_doc,
            context_limit=args.context_limit,
            base_url=base_url,
        )
    elif provider == "anthropic":
        print(f"   → Backend: Anthropic Claude API")
        report = run_comparison_anthropic(
            model=args.model,
            base_doc=base_doc,
            compare_doc=compare_doc,
            context_limit=args.context_limit,
            api_key=args.api_key,
        )
    else:
        print(f"   → Backend: OpenAI-compatible API")
        api_key = args.api_key or os.environ.get("OPENAI_API_KEY", "")
        client_kwargs: dict = {"api_key": api_key}
        if base_url:
            client_kwargs["base_url"] = base_url
        client = OpenAI(**client_kwargs)
        report = run_comparison_openai(
            client=client,
            model=args.model,
            base_doc=base_doc,
            compare_doc=compare_doc,
            context_limit=args.context_limit,
        )

    print(f"\n✅ {report.result.deviation_count} material deviation(s) found")
    if report.total_tokens_used:
        print(f"   → Tokens used: {report.total_tokens_used:,}")

    # ── Write output ─────────────────────────────────────────────────────
    md = render_markdown(report)
    output_path = Path(args.output)
    output_path.write_text(md, encoding="utf-8")
    print(f"📊 Report saved to: {output_path}\n")
    print("=" * 60)
    print(md)


if __name__ == "__main__":
    main()
