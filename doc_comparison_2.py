"""
Document Comparison Tool v2 — Self-RAG Pipeline
================================================
Compares a base PDF against a Word letter of offer using a Self-RAG pipeline:
each clause in the compare document is evaluated against the base document
through a chain of focused LLM decisions before a final deviation judgment
is made and verified.

Pipeline per clause:
    1. ClauseMatchChain      — is this clause substantive enough to look up?
    2. similarity_search     — retrieve top-5 base clauses from InMemoryVectorStore
    3. BestMatchChain        — pick the single best matching base clause
    4. DeviationClassChain   — binary deviation + type (modification/omission/addition/none)
    5. SectionLabelChain  ─┐ — 2-4 word category label          (parallel)
    6. AsyncOpenAI stream ─┘ — stream the explanation comments
    7. GroundingChain        — verify comments are grounded in clause text; retry once if not
    8. MaterialityChain      — legal materiality score 1-5 + severity label

Output table columns:
    # | Section | Base Page | Compare Page | Base Clause | Compare Clause
    | Deviation | Type | Materiality | Comments

Requirements:
    uv add fpdf2 langchain-openai langchain-ollama langchain-core httpx openai pymupdf python-docx

Usage:
    python doc_comparison_2.py \\
        --base files/credit_approval_clauses.pdf \\
        --compare files/letter_of_offer.docx \\
        --base-url http://your-endpoint/v1 \\
        --api-key your-key
"""

import argparse
import asyncio
import os
import tempfile
from pathlib import Path
from typing import Literal, Optional

import fitz  # PyMuPDF
import httpx
from docx import Document as DocxDocument
from fpdf import FPDF
from langchain_core.documents import Document
from langchain_core.prompts import PromptTemplate
from langchain_core.vectorstores import InMemoryVectorStore
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from openai import AsyncOpenAI
from pydantic import BaseModel, Field


# ─────────────────────────────────────────────
# Pydantic models — chain inputs/outputs
# ─────────────────────────────────────────────

class ClauseMatchDecision(BaseModel):
    needs_lookup: bool = Field(
        description=(
            "True if this clause is substantive and warrants comparison against the "
            "base document. False for boilerplate, signature lines, dates, page "
            "headers/footers, or pure formatting text."
        )
    )
    reason: str = Field(description="One-sentence reason for this decision.")


class BestMatchResult(BaseModel):
    best_index: int = Field(
        description=(
            "0-based index of the best matching base-document candidate. "
            "Set to -1 if no candidate is a reasonable match."
        )
    )
    confidence: Literal["High", "Medium", "Low"] = Field(
        description="Confidence in the match quality."
    )


class DeviationClassification(BaseModel):
    has_deviation: bool = Field(
        description=(
            "True if a material deviation exists between the two clauses. "
            "Minor cosmetic rephrasing that preserves identical meaning is NOT a deviation."
        )
    )
    deviation_type: Literal["modification", "omission", "addition", "none"] = Field(
        description=(
            "modification = content or obligation changed; "
            "omission = clause is absent from the compare document; "
            "addition = extra clause in compare document not present in base; "
            "none = no material deviation."
        )
    )


class SectionLabel(BaseModel):
    section: str = Field(
        description=(
            "Concise 2 to 4 word category label for the clause subject matter. "
            "Examples: 'Interest Rate', 'DSCR Covenant', 'Security', 'Repayment Terms'. "
            "Never copy clause text verbatim."
        )
    )


class TargetedContext(BaseModel):
    snippet: str = Field(
        description=(
            "Verbatim excerpt from the base document text that is most directly "
            "relevant to the compare clause. Copy the text exactly as it appears; "
            "do not paraphrase or summarize."
        )
    )


class GroundingVerification(BaseModel):
    grounded: bool = Field(
        description="True if the deviation comment is fully supported by the provided clause texts."
    )
    issues: list[str] = Field(
        default_factory=list,
        description="Specific unsupported or inaccurate claims found in the comment.",
    )


class MaterialityAssessment(BaseModel):
    score: int = Field(
        description=(
            "Materiality score: "
            "5=Critical (fundamental obligation or right changed), "
            "4=Major (significantly alters terms), "
            "3=Minor (noticeable but limited practical impact), "
            "2=Cosmetic+ (slight meaning change), "
            "1=Cosmetic (no practical impact)."
        ),
        ge=1,
        le=5,
    )
    severity: Literal["Critical", "Major", "Minor", "Cosmetic"] = Field(
        description="Severity label matching the score."
    )


# ─────────────────────────────────────────────
# Pydantic models — report output
# ─────────────────────────────────────────────

class DeviationItem(BaseModel):
    item_no: int
    section: str
    base_page: int
    compare_page: int
    base_paragraph: str
    compare_paragraph: str
    deviation: bool
    deviation_type: Literal["modification", "omission", "addition", "none"] = "none"
    score: int = 0
    severity: Literal["Critical", "Major", "Minor", "Cosmetic", "N/A"] = "N/A"
    comments: Optional[str] = None


class ComparisonResult(BaseModel):
    deviations: list[DeviationItem]

    @property
    def deviation_count(self) -> int:
        return sum(1 for d in self.deviations if d.deviation)


class ComparisonReport(BaseModel):
    base_doc: str
    compare_doc: str
    result: ComparisonResult


# ─────────────────────────────────────────────
# DOCX XML page-break helpers
# ─────────────────────────────────────────────

# XML namespace identifier — not a network URL; baked into every .docx file
_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _page_breaks_in_paragraph(para) -> int:
    """
    Count page breaks inside a paragraph's XML, covering:
    - w:lastRenderedPageBreak  (soft/auto breaks inserted by Word on save)
    - w:br[@w:type='page']     (explicit hard breaks via Ctrl+Enter)
    """
    count = 0
    for elem in para._p.iter():
        local = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if local == "lastRenderedPageBreak":
            count += 1
        elif local == "br" and elem.get(f"{{{_W}}}type") == "page":
            count += 1
    return count



# ─────────────────────────────────────────────
# DOCX → PDF conversion (fpdf2, no external tools)
# ─────────────────────────────────────────────

def _extract_docx_paragraphs(docx_path: str) -> list[tuple[str, int]]:
    """Return (text, page_number) pairs using XML page-break detection."""
    result: list[tuple[str, int]] = []
    page = 1
    for para in DocxDocument(docx_path).paragraphs:
        page += _page_breaks_in_paragraph(para)
        text = para.text.strip()
        if text:
            result.append((text, page))
    return result


def convert_docx_to_pdf(docx_path: str) -> str:
    """
    Convert a DOCX to a temporary PDF using fpdf2.
    Page boundaries follow w:lastRenderedPageBreak XML markers.
    Returns the temp PDF path — caller is responsible for deletion.
    """
    paragraphs = _extract_docx_paragraphs(docx_path)

    pdf = FPDF()
    pdf.set_margins(15, 15, 15)
    pdf.set_auto_page_break(auto=False)
    pdf.add_page()
    current_page = 1

    for text, page_no in paragraphs:
        while current_page < page_no:
            pdf.add_page()
            current_page += 1
        pdf.set_font("Helvetica", size=10)
        # fpdf2 uses latin-1 internally; replace unmappable characters
        safe = text.encode("latin-1", errors="replace").decode("latin-1")
        pdf.multi_cell(0, 6, safe)
        pdf.ln(2)

    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    tmp.close()
    pdf.output(tmp.name)
    return tmp.name


# ─────────────────────────────────────────────
# Document loading
# ─────────────────────────────────────────────

def load_pdf_as_documents(path: str, doc_name: Optional[str] = None) -> list[Document]:
    """
    Extract text from a PDF with PyMuPDF, aggregated per page.
    Each page produces one Document containing all its text blocks joined by newlines.
    Returns LangChain Documents with page and doc_name metadata.
    """
    name = doc_name or Path(path).name
    docs: list[Document] = []

    with fitz.open(path) as pdf:
        for page_no, page in enumerate(pdf, start=1):
            blocks = [
                block[4].strip()
                for block in page.get_text("blocks")
                if block[4].strip()
            ]
            if blocks:
                docs.append(Document(
                    page_content="\n".join(blocks),
                    metadata={"page": page_no, "doc_name": name},
                ))

    return docs


# ─────────────────────────────────────────────
# Chain definitions (mirrors temp.py pattern)
# ─────────────────────────────────────────────

def build_chains(llm: ChatOpenAI) -> dict:
    """Build all structured-output LangChain chains."""

    clause_match_chain = PromptTemplate(
        input_variables=["clause"],
        template=(
            "You are reviewing a clause from a legal offer letter.\n\n"
            "Clause:\n{clause}\n\n"
            "Decide whether this clause is substantive enough to compare against "
            "a base credit approval document. Respond False for boilerplate, "
            "signature lines, dates, page headers/footers, or pure formatting."
        ),
    ) | llm.with_structured_output(ClauseMatchDecision)

    best_match_chain = PromptTemplate(
        input_variables=["compare_clause", "candidates"],
        template=(
            "You are matching a compare-document clause to base-document candidates.\n\n"
            "COMPARE CLAUSE:\n{compare_clause}\n\n"
            "BASE CANDIDATES (0-indexed):\n{candidates}\n\n"
            "Select the index of the single best matching candidate. "
            "Set best_index=-1 and confidence='Low' if none are a reasonable match."
        ),
    ) | llm.with_structured_output(BestMatchResult)

    deviation_classification_chain = PromptTemplate(
        input_variables=["base_clause", "compare_clause"],
        template=(
            "You are a senior legal reviewer.\n\n"
            "BASE CLAUSE (reference standard):\n{base_clause}\n\n"
            "COMPARE CLAUSE (document under review):\n{compare_clause}\n\n"
            "Determine whether a material deviation exists. "
            "Classify as: modification (content changed), omission (absent from compare), "
            "addition (extra in compare), or none. "
            "Minor cosmetic rephrasing that preserves identical meaning is NOT a deviation."
        ),
    ) | llm.with_structured_output(DeviationClassification)

    section_label_chain = PromptTemplate(
        input_variables=["base_clause", "compare_clause"],
        template=(
            "Provide a concise 2 to 4 word category label for the subject matter "
            "of these two clauses.\n"
            "Examples: 'Interest Rate', 'DSCR Covenant', 'Security', 'Repayment Terms'.\n"
            "Do NOT copy text from the clauses.\n\n"
            "Base clause: {base_clause}\n"
            "Compare clause: {compare_clause}"
        ),
    ) | llm.with_structured_output(SectionLabel)

    targeted_context_chain = PromptTemplate(
        input_variables=["base_text", "compare_clause"],
        template=(
            "From the base document text below, extract the verbatim excerpt that is "
            "most directly relevant to the compare clause.\n\n"
            "BASE TEXT:\n{base_text}\n\n"
            "COMPARE CLAUSE:\n{compare_clause}\n\n"
            "Copy the relevant excerpt exactly as it appears in the base text. "
            "Do not paraphrase, summarize, or add any words."
        ),
    ) | llm.with_structured_output(TargetedContext)

    grounding_chain = PromptTemplate(
        input_variables=["comments", "base_clause", "compare_clause"],
        template=(
            "Verify whether this deviation comment is fully supported by the two "
            "clause texts provided.\n\n"
            "COMMENT:\n{comments}\n\n"
            "BASE CLAUSE:\n{base_clause}\n\n"
            "COMPARE CLAUSE:\n{compare_clause}\n\n"
            "List any specific claims in the comment that are not directly evidenced "
            "by the clause texts."
        ),
    ) | llm.with_structured_output(GroundingVerification)

    materiality_chain = PromptTemplate(
        input_variables=["deviation_type", "comments", "base_clause", "compare_clause"],
        template=(
            "Rate the legal/contractual materiality of this deviation.\n\n"
            "Deviation type: {deviation_type}\n"
            "Comment: {comments}\n"
            "Base clause: {base_clause}\n"
            "Compare clause: {compare_clause}\n\n"
            "Score 1 to 5:\n"
            "  5 = Critical (fundamental obligation or right changed)\n"
            "  4 = Major (significantly alters terms)\n"
            "  3 = Minor (noticeable but limited practical impact)\n"
            "  2 = Cosmetic+ (slight meaning change)\n"
            "  1 = Cosmetic (no practical impact)\n"
            "Assign severity: Critical / Major / Minor / Cosmetic."
        ),
    ) | llm.with_structured_output(MaterialityAssessment)

    return {
        "clause_match": clause_match_chain,
        "best_match": best_match_chain,
        "targeted_context": targeted_context_chain,
        "deviation_classification": deviation_classification_chain,
        "section_label": section_label_chain,
        "grounding": grounding_chain,
        "materiality": materiality_chain,
    }


# ─────────────────────────────────────────────
# Streaming generation (Step 6)
# ─────────────────────────────────────────────

async def stream_deviation_comments(
    openai_client: AsyncOpenAI,
    model: str,
    base_clause: str,
    compare_clause: str,
    deviation_type: str,
    correction_notes: Optional[list[str]] = None,
) -> str:
    """Stream a concise one-sentence deviation explanation from the LLM."""
    correction_block = ""
    if correction_notes:
        items = "\n".join(f"  - {note}" for note in correction_notes)
        correction_block = (
            f"\n\nA prior attempt contained unsupported claims. "
            f"Correct these specific issues:\n{items}"
        )

    prompt = (
        f"You are a senior legal reviewer. Write one concise sentence (max 40 words) "
        f"explaining the deviation between these two clauses.\n\n"
        f"Deviation type: {deviation_type}\n"
        f"Base clause: {base_clause}\n"
        f"Compare clause: {compare_clause}"
        f"{correction_block}\n\n"
        f"Write only the explanation sentence, nothing else."
    )

    chunks: list[str] = []
    stream = await openai_client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        temperature=0,
        stream=True,
    )
    async for chunk in stream:
        delta = chunk.choices[0].delta.content or ""
        chunks.append(delta)

    return "".join(chunks).strip()


# ─────────────────────────────────────────────
# Self-RAG per-clause comparison
# ─────────────────────────────────────────────

async def compare_clause(
    compare_doc: Document,
    item_no: int,
    vectorstore: InMemoryVectorStore,
    openai_client: AsyncOpenAI,
    model: str,
    chains: dict,
    semaphore: asyncio.Semaphore,
) -> DeviationItem:
    async with semaphore:
        compare_text = compare_doc.page_content
        compare_page = compare_doc.metadata.get("page", 0)

        def _no_deviation(section: str = "N/A", base_text: str = "", base_pg: int = 0) -> DeviationItem:
            return DeviationItem(
                item_no=item_no,
                section=section,
                base_page=base_pg,
                compare_page=compare_page,
                base_paragraph=base_text,
                compare_paragraph=compare_text,
                deviation=False,
                deviation_type="none",
                score=0,
                severity="N/A",
                comments=None,
            )

        # ── Step 1: ClauseMatchChain ──────────────────────────────────────
        try:
            match_decision: ClauseMatchDecision = await chains["clause_match"].ainvoke(
                {"clause": compare_text}
            )
        except Exception:
            return _no_deviation()

        if not match_decision.needs_lookup:
            return _no_deviation()

        # ── Step 2: Retrieve top-5 from base doc ──────────────────────────
        retrieved = await vectorstore.asimilarity_search(compare_text, k=5)
        if not retrieved:
            return _no_deviation()

        # ── Step 3: BestMatchChain ────────────────────────────────────────
        candidates_text = "\n\n".join(
            f"[{i}] (p.{doc.metadata.get('page', '?')}): {doc.page_content}"
            for i, doc in enumerate(retrieved)
        )
        try:
            best: BestMatchResult = await chains["best_match"].ainvoke({
                "compare_clause": compare_text,
                "candidates": candidates_text,
            })
        except Exception:
            return _no_deviation()

        if best.confidence == "Low" or not (0 <= best.best_index < len(retrieved)):
            return _no_deviation()

        base_doc = retrieved[best.best_index]
        base_text = base_doc.page_content
        base_page = base_doc.metadata.get("page", 0)

        # ── Step 3.5: TargetedContextChain ───────────────────────────────
        # Narrow the full-page base text down to the specific verbatim snippet
        # most relevant to the compare clause; all downstream chains use this.
        try:
            targeted: TargetedContext = await chains["targeted_context"].ainvoke({
                "base_text": base_text,
                "compare_clause": compare_text,
            })
            base_snippet = targeted.snippet.strip() or base_text
        except Exception:
            base_snippet = base_text

        # ── Step 4: DeviationClassificationChain ─────────────────────────
        try:
            dev_class: DeviationClassification = await chains["deviation_classification"].ainvoke({
                "base_clause": base_snippet,
                "compare_clause": compare_text,
            })
        except Exception:
            return _no_deviation(base_text=base_snippet, base_pg=base_page)

        if not dev_class.has_deviation:
            try:
                sec: SectionLabel = await chains["section_label"].ainvoke({
                    "base_clause": base_snippet,
                    "compare_clause": compare_text,
                })
                section = sec.section
            except Exception:
                section = "N/A"
            return _no_deviation(section=section, base_text=base_snippet, base_pg=base_page)

        # ── Steps 5 & 6: SectionLabelChain (parallel) + streaming comments
        section_task = asyncio.create_task(
            chains["section_label"].ainvoke({
                "base_clause": base_snippet,
                "compare_clause": compare_text,
            })
        )

        comments = await stream_deviation_comments(
            openai_client=openai_client,
            model=model,
            base_clause=base_snippet,
            compare_clause=compare_text,
            deviation_type=dev_class.deviation_type,
        )

        try:
            sec_result: SectionLabel = await section_task
            section = sec_result.section
        except Exception:
            section = "N/A"

        # ── Step 7: GroundingChain — verify; retry once if weak ───────────
        try:
            grounding: GroundingVerification = await chains["grounding"].ainvoke({
                "comments": comments,
                "base_clause": base_snippet,
                "compare_clause": compare_text,
            })
            if not grounding.grounded and grounding.issues:
                comments = await stream_deviation_comments(
                    openai_client=openai_client,
                    model=model,
                    base_clause=base_snippet,
                    compare_clause=compare_text,
                    deviation_type=dev_class.deviation_type,
                    correction_notes=grounding.issues,
                )
        except Exception:
            pass  # keep original comments if grounding check fails

        # ── Step 8: MaterialityChain ──────────────────────────────────────
        try:
            materiality: MaterialityAssessment = await chains["materiality"].ainvoke({
                "deviation_type": dev_class.deviation_type,
                "comments": comments,
                "base_clause": base_snippet,
                "compare_clause": compare_text,
            })
            score = materiality.score
            severity = materiality.severity
        except Exception:
            score = 0
            severity = "N/A"

        print(f"   [clause {item_no}] ✓ deviation={dev_class.deviation_type} | {section} | materiality={score}")

        return DeviationItem(
            item_no=item_no,
            section=section,
            base_page=base_page,
            compare_page=compare_page,
            base_paragraph=base_snippet,
            compare_paragraph=compare_text,
            deviation=True,
            deviation_type=dev_class.deviation_type,
            score=score,
            severity=severity,
            comments=comments,
        )


# ─────────────────────────────────────────────
# Markdown renderer
# ─────────────────────────────────────────────

def _truncate(text: str, max_chars: int = 250) -> str:
    text = text.replace("\n", " ").strip()
    return text if len(text) <= max_chars else text[:max_chars].rstrip() + "…"


def render_markdown(report: ComparisonReport) -> str:
    all_items = report.result.deviations

    if not all_items:
        body = "_No clauses examined._\n"
    else:
        header = (
            f"| # | Section "
            f"| `{report.base_doc}` Page "
            f"| `{report.compare_doc}` Page "
            f"| Base Clause "
            f"| Compare Clause "
            f"| Deviation "
            f"| Type "
            f"| Materiality "
            f"| Comments |\n"
            "|---|---|---|---|---|---|---|---|---|---|\n"
        )
        rows = []
        for d in all_items:
            base_pg = f"p.{d.base_page}" if d.base_page else "—"
            cmp_pg = f"p.{d.compare_page}" if d.compare_page else "—"
            materiality = f"{d.score}/5 ({d.severity})" if d.score else "—"
            rows.append(
                f"| {d.item_no} "
                f"| {d.section} "
                f"| {base_pg} "
                f"| {cmp_pg} "
                f"| {_truncate(d.base_paragraph)} "
                f"| {_truncate(d.compare_paragraph)} "
                f"| {'Yes' if d.deviation else 'No'} "
                f"| {d.deviation_type} "
                f"| {materiality} "
                f"| {d.comments or '—'} |\n"
            )
        body = header + "".join(rows)

    return (
        f"# Document Comparison Report (Self-RAG)\n\n"
        f"**Base document:** `{report.base_doc}`  \n"
        f"**Compared document:** `{report.compare_doc}`  \n"
        f"**Material deviations found:** {report.result.deviation_count}  \n"
        f"\n---\n\n"
        f"## Deviation Table\n\n"
        f"{body}\n"
        f"---\n"
        f"*Generated by doc_comparison_2.py — Self-RAG pipeline*\n"
    )


# ─────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────

async def main() -> None:
    parser = argparse.ArgumentParser(
        description="Document comparison — Self-RAG pipeline with InMemoryVectorStore."
    )
    parser.add_argument("--base",        required=True,  help="Path to base PDF document")
    parser.add_argument("--compare",     required=True,  help="Path to compare DOCX document")
    parser.add_argument("--output",      default="comparison_report_2.md", help="Output Markdown file")
    parser.add_argument("--model",       default="gpt-oss-120b", help="LLM model name")
    parser.add_argument("--base-url",    required=True,  help="API base URL (LLM + embeddings)")
    parser.add_argument("--api-key",     default=None,   help="API key (falls back to OPENAI_API_KEY)")
    parser.add_argument("--embed-model", default="bge-m3", help="Embedding model name")
    parser.add_argument("--concurrency", default=5, type=int, help="Max parallel clause comparisons")
    args = parser.parse_args()

    api_key = args.api_key or os.environ.get("OPENAI_API_KEY", "no-key")

    # httpx.AsyncClient provides shared transport: custom headers + timeout.
    # base_url is NOT set here — AsyncOpenAI and ChatOpenAI each manage their own.
    async with httpx.AsyncClient(
        headers={"api-key": api_key},
        timeout=httpx.Timeout(120.0),
    ) as httpx_client:

        # ── Clients ──────────────────────────────────────────────────────
        openai_client = AsyncOpenAI(
            base_url=args.base_url,
            api_key=api_key,
            http_client=httpx_client,
        )

        langchain_llm = ChatOpenAI(
            openai_api_base=args.base_url,
            openai_api_key=api_key,
            model_name=args.model,
            temperature=0,
            streaming=True,
        )

        embeddings = OpenAIEmbeddings(
            openai_api_base=args.base_url,
            openai_api_key=api_key,
            model=args.embed_model,
        )

        # ── Chains ───────────────────────────────────────────────────────
        chains = build_chains(langchain_llm)

        # ── Load documents ───────────────────────────────────────────────
        print(f"📄 Loading base PDF: {args.base}")
        base_docs = load_pdf_as_documents(args.base)
        print(f"   → {len(base_docs)} chunks extracted")

        print(f"📝 Converting and loading compare DOCX: {args.compare}")
        tmp_pdf = convert_docx_to_pdf(args.compare)
        try:
            compare_docs = load_pdf_as_documents(tmp_pdf, doc_name=Path(args.compare).name)
            print(f"   → {len(compare_docs)} chunks extracted")
        finally:
            Path(tmp_pdf).unlink(missing_ok=True)

        # ── Build vector store ───────────────────────────────────────────
        print("\n🔍 Indexing base document into InMemoryVectorStore...")
        vectorstore = InMemoryVectorStore(embedding=embeddings)
        await vectorstore.aadd_documents(base_docs)
        print(f"   → {len(base_docs)} chunks indexed")

        # ── Run Self-RAG comparison ──────────────────────────────────────
        print(
            f"\n🤖 Running Self-RAG comparison "
            f"({len(compare_docs)} clauses, concurrency={args.concurrency})..."
        )
        semaphore = asyncio.Semaphore(args.concurrency)

        tasks = [
            compare_clause(
                compare_doc=doc,
                item_no=i + 1,
                vectorstore=vectorstore,
                openai_client=openai_client,
                model=args.model,
                chains=chains,
                semaphore=semaphore,
            )
            for i, doc in enumerate(compare_docs)
        ]
        items: list[DeviationItem] = list(await asyncio.gather(*tasks))

        # ── Build report ─────────────────────────────────────────────────
        result = ComparisonResult(deviations=items)
        report = ComparisonReport(
            base_doc=Path(args.base).name,
            compare_doc=Path(args.compare).name,
            result=result,
        )

        print(f"\n✅ {result.deviation_count} material deviation(s) found")

        # ── Write output ─────────────────────────────────────────────────
        md = render_markdown(report)
        output_path = Path(args.output)
        output_path.write_text(md, encoding="utf-8")
        print(f"📊 Report saved to: {output_path}\n")
        print("=" * 60)
        print(md)


if __name__ == "__main__":
    asyncio.run(main())
